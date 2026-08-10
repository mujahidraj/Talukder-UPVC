import sys
import pandas as pd
from docx import Document

def extract_docx(file_path):
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip().replace('\n', ' '))
            text.append(" | ".join(row_data))
    return "\n".join(text)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: extract.py <docx/xlsx>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    if file_path.endswith('.docx'):
        try:
            with open("docx_content.txt", "w", encoding="utf-8") as f:
                f.write(extract_docx(file_path))
            print("Successfully extracted docx to docx_content.txt")
        except Exception as e:
            print(f"Error reading docx: {e}")
    elif file_path.endswith('.xlsx'):
        try:
            with open("xlsx_content.txt", "w", encoding="utf-8") as f:
                df = pd.read_excel(file_path, nrows=5)
                f.write("Columns: " + str(df.columns.tolist()) + "\n")
                f.write("First row sample:\n" + str(df.head(1).to_dict(orient='records')))
            print("Successfully extracted xlsx to xlsx_content.txt")
        except Exception as e:
            print(f"Error reading xlsx: {e}")
